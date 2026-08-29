from typing import Dict, Any
import os
from datetime import datetime
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
from app.core.config import settings


class PDFExporter:
    def __init__(self):
        self.template_env = Environment(
            loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), "templates"))
        )
    
    def export(self, plan_data: Dict[str, Any], output_path: str) -> str:
        template = self.template_env.get_template("plan_pdf.html")
        html_content = template.render(
            plan=plan_data,
            generated_at=datetime.now().strftime("%B %d, %Y")
        )
        
        HTML(string=html_content).write_pdf(output_path)
        return output_path


class DOCXExporter:
    def __init__(self):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self.Document = Document
            self.Inches = Inches
            self.Pt = Pt
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed")
    
    def export(self, plan_data: Dict[str, Any], output_path: str) -> str:
        doc = self.Document()
        
        # Title
        title = doc.add_heading(plan_data.get("plan", {}).get("name", "Business Plan"), 0)
        title.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{plan_data.get('plan', {}).get('frequency', 'Quarterly').title()} Plan • Generated {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = self.Pt(12)
        run.font.color.rgb = self.RGBColor(128, 128, 128)
        
        # Executive Summary
        self._add_section(doc, "Executive Summary", plan_data.get("generated_plan", ""))
        
        # Market Analysis
        market = plan_data.get("market_analysis", {})
        if market:
            doc.add_heading("Market Analysis", level=1)
            self._add_key_value(doc, "TAM", f"${market.get('tam', 0):,.0f}")
            self._add_key_value(doc, "SAM", f"${market.get('sam', 0):,.0f}")
            self._add_key_value(doc, "SOM", f"${market.get('som', 0):,.0f}")
            self._add_key_value(doc, "Market Growth Rate", f"{market.get('market_growth_rate', 0)*100:.1f}%")
            
            if market.get("key_trends"):
                doc.add_heading("Key Trends", level=2)
                for trend in market["key_trends"]:
                    doc.add_paragraph(trend, style="List Bullet")
        
        # Financial Projections
        financial = plan_data.get("financial_projections", {})
        if financial:
            doc.add_heading("Financial Projections", level=1)
            assumptions = financial.get("assumptions", {})
            doc.add_heading("Key Assumptions", level=2)
            for key, value in assumptions.items():
                if isinstance(value, float):
                    self._add_key_value(doc, key.replace("_", " ").title(), f"{value*100:.1f}%")
                else:
                    self._add_key_value(doc, key.replace("_", " ").title(), str(value))
            
            # Add summary table
            pnl = financial.get("pnl", [])
            if pnl:
                doc.add_heading("P&L Summary (Annual)", level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                headers = ["Year", "Revenue", "Gross Profit", "EBITDA", "Net Income"]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                
                # Aggregate by year
                yearly = {}
                for row in pnl:
                    period = row.get("period", "")
                    if "Year" in period or "Month 12" in period or "Month 24" in period or "Month 36" in period:
                        year = period.split()[-1]
                        if year not in yearly:
                            yearly[year] = {"revenue": 0, "gross_profit": 0, "ebitda": 0, "net_income": 0}
                        yearly[year]["revenue"] += row.get("revenue", 0)
                        yearly[year]["gross_profit"] += row.get("gross_profit", 0)
                        yearly[year]["ebitda"] += row.get("ebitda", 0)
                        yearly[year]["net_income"] += row.get("net_income", 0)
                
                for year, vals in sorted(yearly.items()):
                    row = table.add_row()
                    row.cells[0].text = year
                    row.cells[1].text = f"${vals['revenue']:,.0f}"
                    row.cells[2].text = f"${vals['gross_profit']:,.0f}"
                    row.cells[3].text = f"${vals['ebitda']:,.0f}"
                    row.cells[4].text = f"${vals['net_income']:,.0f}"
        
        # Strategy
        strategy = plan_data.get("strategy", {})
        if strategy:
            doc.add_heading("Strategy", level=1)
            
            swot = strategy.get("swot", {})
            if swot:
                doc.add_heading("SWOT Analysis", level=2)
                for quadrant, items in swot.items():
                    doc.add_heading(quadrant.title(), level=3)
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")
            
            okrs = strategy.get("okrs", [])
            if okrs:
                doc.add_heading("OKRs", level=2)
                for okr in okrs:
                    doc.add_heading(okr.get("objective", ""), level=3)
                    for kr in okr.get("key_results", []):
                        doc.add_paragraph(f"{kr.get('metric', '')}: {kr.get('target', 0)} {kr.get('unit', '')}", style="List Bullet")
        
        # Competitors
        competitors = plan_data.get("competitor_analysis", {}).get("competitors", [])
        if competitors:
            doc.add_heading("Competitive Analysis", level=1)
            for comp in competitors:
                doc.add_heading(comp.get("name", ""), level=2)
                self._add_key_value(doc, "Website", comp.get("website", ""))
                self._add_key_value(doc, "Pricing Model", comp.get("pricing_model", ""))
                if comp.get("key_features"):
                    doc.add_paragraph("Key Features:", style="List Bullet")
                    for feat in comp["key_features"][:5]:
                        doc.add_paragraph(feat, style="List Bullet 2")
        
        doc.save(output_path)
        return output_path
    
    def _add_section(self, doc, heading: str, content: str):
        doc.add_heading(heading, level=1)
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    
    def _add_key_value(self, doc, key: str, value: str):
        p = doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        run_key.bold = True
        p.add_run(value)


class XLSXExporter:
    def __init__(self):
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            self.openpyxl = openpyxl
            self.Font = Font
            self.PatternFill = PatternFill
            self.Alignment = Alignment
            self.Border = Border
            self.Side = Side
            self.get_column_letter = get_column_letter
        except ImportError:
            raise ImportError("openpyxl not installed")
    
    def export(self, plan_data: Dict[str, Any], output_path: str) -> str:
        wb = self.openpyxl.Workbook()
        
        # Sheet 1: Summary
        ws1 = wb.active
        ws1.title = "Summary"
        self._write_summary_sheet(ws1, plan_data)
        
        # Sheet 2: Financial Projections
        ws2 = wb.create_sheet("Financials")
        self._write_financial_sheet(ws2, plan_data.get("financial_projections", {}))
        
        # Sheet 3: Market Analysis
        ws3 = wb.create_sheet("Market Analysis")
        self._write_market_sheet(ws3, plan_data.get("market_analysis", {}))
        
        # Sheet 4: Competitors
        ws4 = wb.create_sheet("Competitors")
        self._write_competitor_sheet(ws4, plan_data.get("competitor_analysis", {}).get("competitors", []))
        
        # Sheet 5: OKRs & Milestones
        ws5 = wb.create_sheet("OKRs & Milestones")
        self._write_okr_sheet(ws5, plan_data.get("strategy", {}))
        
        wb.save(output_path)
        return output_path
    
    def _write_summary_sheet(self, ws, plan_data: Dict[str, Any]):
        plan = plan_data.get("plan", {})
        headers = ["Field", "Value"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = self.Font(bold=True)
            cell.fill = self.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.font = self.Font(bold=True, color="FFFFFF")
        
        data = [
            ("Plan Name", plan.get("name", "")),
            ("Frequency", plan.get("frequency", "").title()),
            ("Industry", plan.get("industry", "")),
            ("Company Size", plan.get("company_size", "")),
            ("Revenue Range", plan.get("revenue_range", "")),
            ("Status", plan.get("status", "").title()),
            ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ]
        
        for row, (key, value) in enumerate(data, 2):
            ws.cell(row=row, column=1, value=key)
            ws.cell(row=row, column=2, value=value)
        
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 50
    
    def _write_financial_sheet(self, ws, financial: Dict[str, Any]):
        pnl = financial.get("pnl", [])
        if not pnl:
            ws.cell(row=1, column=1, value="No financial data available")
            return
        
        headers = ["Period", "Revenue", "COGS", "Gross Profit", "OpEx", "EBITDA", "Depreciation", "Interest", "Tax", "Net Income"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = self.Font(bold=True)
            cell.fill = self.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.font = self.Font(bold=True, color="FFFFFF")
        
        for row, item in enumerate(pnl, 2):
            ws.cell(row=row, column=1, value=item.get("period", ""))
            ws.cell(row=row, column=2, value=item.get("revenue", 0))
            ws.cell(row=row, column=3, value=item.get("cogs", 0))
            ws.cell(row=row, column=4, value=item.get("gross_profit", 0))
            ws.cell(row=row, column=5, value=item.get("operating_expenses", 0))
            ws.cell(row=row, column=6, value=item.get("ebitda", 0))
            ws.cell(row=row, column=7, value=item.get("depreciation", 0))
            ws.cell(row=row, column=8, value=item.get("interest", 0))
            ws.cell(row=row, column=9, value=item.get("tax", 0))
            ws.cell(row=row, column=10, value=item.get("net_income", 0))
        
        # Format currency columns
        for col in range(2, 11):
            for row in range(2, len(pnl) + 2):
                ws.cell(row=row, column=col).number_format = '$#,##0.00'
        
        for col in range(1, 11):
            ws.column_dimensions[self.get_column_letter(col)].width = 18
    
    def _write_market_sheet(self, ws, market: Dict[str, Any]):
        if not market:
            ws.cell(row=1, column=1, value="No market data available")
            return
        
        row = 1
        for key, value in market.items():
            if isinstance(value, (list, dict)):
                continue
            ws.cell(row=row, column=1, value=key.replace("_", " ").title()).font = self.Font(bold=True)
            ws.cell(row=row, column=2, value=value)
            row += 1
        
        if market.get("key_trends"):
            ws.cell(row=row, column=1, value="Key Trends").font = self.Font(bold=True)
            row += 1
            for trend in market["key_trends"]:
                ws.cell(row=row, column=2, value=trend)
                row += 1
        
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 60
    
    def _write_competitor_sheet(self, ws, competitors: list):
        if not competitors:
            ws.cell(row=1, column=1, value="No competitor data available")
            return
        
        headers = ["Name", "Website", "Pricing Model", "Key Features", "Tech Stack", "Positioning", "Strengths", "Weaknesses"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = self.Font(bold=True)
            cell.fill = self.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.font = self.Font(bold=True, color="FFFFFF")
        
        for row, comp in enumerate(competitors, 2):
            ws.cell(row=row, column=1, value=comp.get("name", ""))
            ws.cell(row=row, column=2, value=comp.get("website", ""))
            ws.cell(row=row, column=3, value=comp.get("pricing_model", ""))
            ws.cell(row=row, column=4, value="; ".join(comp.get("key_features", [])[:5]))
            ws.cell(row=row, column=5, value="; ".join(comp.get("tech_stack", [])))
            ws.cell(row=row, column=6, value=comp.get("positioning", ""))
            ws.cell(row=row, column=7, value="; ".join(comp.get("strengths", [])))
            ws.cell(row=row, column=8, value="; ".join(comp.get("weaknesses", [])))
        
        for col in range(1, 9):
            ws.column_dimensions[self.get_column_letter(col)].width = 25
    
    def _write_okr_sheet(self, ws, strategy: Dict[str, Any]):
        okrs = strategy.get("okrs", [])
        milestones = strategy.get("milestones", [])
        
        if okrs:
            ws.cell(row=1, column=1, value="OKRs").font = self.Font(bold=True, size=14)
            row = 2
            headers = ["Objective", "Key Result", "Target", "Current", "Unit", "Owner", "Timeline"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = self.Font(bold=True)
                cell.fill = self.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.font = self.Font(bold=True, color="FFFFFF")
            
            for okr in okrs:
                for kr in okr.get("key_results", []):
                    row += 1
                    ws.cell(row=row, column=1, value=okr.get("objective", ""))
                    ws.cell(row=row, column=2, value=kr.get("metric", ""))
                    ws.cell(row=row, column=3, value=kr.get("target", 0))
                    ws.cell(row=row, column=4, value=kr.get("current", 0))
                    ws.cell(row=row, column=5, value=kr.get("unit", ""))
                    ws.cell(row=row, column=6, value=okr.get("owner", ""))
                    ws.cell(row=row, column=7, value=okr.get("timeline", ""))
        
        if milestones:
            start_row = row + 3
            ws.cell(row=start_row, column=1, value="Milestones").font = self.Font(bold=True, size=14)
            row = start_row + 1
            headers = ["Milestone", "Target Date", "Dependencies", "Success Criteria"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = self.Font(bold=True)
                cell.fill = self.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.font = self.Font(bold=True, color="FFFFFF")
            
            for ms in milestones:
                row += 1
                ws.cell(row=row, column=1, value=ms.get("milestone", ""))
                ws.cell(row=row, column=2, value=ms.get("target_date", ""))
                ws.cell(row=row, column=3, value="; ".join(ms.get("dependencies", [])))
                ws.cell(row=row, column=4, value=ms.get("success_criteria", ""))